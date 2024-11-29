from PyPDF2 import PdfReader, PdfWriter
import docx

# PDF: Extract text from a PDF file
reader = PdfReader("example.pdf")
page = reader.pages[0]
text = page.extract_text()
print(f"PDF Text: {text}")

# PDF: Merge two PDFs
pdf_writer = PdfWriter()
pdf_writer.add_page(reader.pages[0])
with open("merged.pdf", "wb") as output_pdf:
    pdf_writer.write(output_pdf)

# Word: Create a Word document and add text
doc = docx.Document()
doc.add_heading("Sample Word Document", 0)
doc.add_paragraph("This is a paragraph in the document.")
doc.save("example.docx")

# Word: Read a Word document
doc = docx.Document("example.docx")
for para in doc.paragraphs:
    print(para.text)